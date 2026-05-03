"""
Extract exam questions from PDF and DOCX files for Nhật Khôi and Minh Khanh.
"""
import pdfplumber
import docx
import os

BASE = r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem"

def read_pdf(path):
    """Extract all text from a PDF file."""
    text = ""
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception as e:
        text = f"[ERROR reading {path}: {e}]"
    return text

def read_docx(path):
    """Extract all text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(path)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        text = f"[ERROR reading {path}: {e}]"
    return text

# Files to read
files = {
    # Folder 1: 10 de kiem tra cuoi HK2 Toan 8
    "CK2_T8_DS1": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\10-de-kiem-tra-cuoi-hoc-ki-2-toan-8-knttvcs-nam-hoc-2023-2024\01.CK-2-T-8-KNTT-DS-1.docx",
    "CK2_T8_DS2": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\10-de-kiem-tra-cuoi-hoc-ki-2-toan-8-knttvcs-nam-hoc-2023-2024\01.CK-2-T-8-KNTT-DS-2.docx",
    "CK2_T8_DS3": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\10-de-kiem-tra-cuoi-hoc-ki-2-toan-8-knttvcs-nam-hoc-2023-2024\01.CK-2-T-8-KNTT-DS-3.docx",
    "CK2_T8_DS4": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\10-de-kiem-tra-cuoi-hoc-ki-2-toan-8-knttvcs-nam-hoc-2023-2024\01.CK-2-T-8-KNTT-DS-4.docx",
    "CK2_T8_DS5": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\10-de-kiem-tra-cuoi-hoc-ki-2-toan-8-knttvcs-nam-hoc-2023-2024\01.CK-2-T-8-KNTT-DS-5.docx",
    # Folder 2: Khoi 8 HK2 PDFs
    "HK2_HaiDuong": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Khoi 8 HK2\de-cuoi-hoc-ki-2-toan-8-nam-2024-2025-phong-gddt-tp-hai-duong.pdf",
    "HK2_LeHongPhong": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Khoi 8 HK2\de-cuoi-hoc-ki-2-toan-8-nam-2024-2025-truong-le-hong-phong-quang-nam.pdf",
    "HK2_LySon": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Khoi 8 HK2\de-cuoi-hoc-ki-2-toan-8-nam-2024-2025-truong-thcs-ly-son-ha-noi.pdf",
    "HK2_VoTruongToan": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Khoi 8 HK2\de-cuoi-hoc-ki-2-toan-8-nam-2024-2025-truong-thcs-vo-truong-toan-br-vt.pdf",
    "HK2_TrucNinh": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Khoi 8 HK2\de-hoc-ky-2-toan-8-nam-2024-2025-phong-gddt-truc-ninh-nam-dinh.pdf",
    "HK2_BacGiang": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Khoi 8 HK2\de-kiem-tra-cuoi-hoc-ki-2-toan-8-nam-2024-2025-so-gddt-bac-giang.pdf",
    # Folder 3: Lop 7 - Tieng Anh
    "TA7_1_docx": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Nhat Khoi lop 7 - Ôn Tập Lại\Khoi\TA 7 - 1.docx",
    "TA7_2_pdf": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Nhat Khoi lop 7 - Ôn Tập Lại\Khoi\ta 7 - 2.pdf",
    "TEST_VIC_TA_K8": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Nhat Khoi lop 7 - Ôn Tập Lại\Khoi\[TEST] VIC_TA_HK1_K8.pdf",
    "De1_doc": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Nhat Khoi lop 7 - Ôn Tập Lại\Khoi\De 1.doc",
    "De2_docx": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Nhat Khoi lop 7 - Ôn Tập Lại\Khoi\De 2.docx",
    "Doc1_docx": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Nhat Khoi lop 7 - Ôn Tập Lại\Khoi\Doc1.docx",
    "In0409_docx": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Nhat Khoi lop 7 - Ôn Tập Lại\Khoi\In 04 09.docx",
    "0710_docx": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Nhật Khôi - 8A1\Nhat Khoi lop 7 - Ôn Tập Lại\Khoi\07 10.docx",
    # Minh Khanh - English PDFs
    "MK_De04_Lop6": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Minh Khanh - 6A3\De 04 - Lop 6.pdf",
    "MK_De05_Lop6": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Minh Khanh - 6A3\De 05 - Lop 6.pdf",
    "MK_De06_Lop6": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Minh Khanh - 6A3\De 06 - Lop 6.pdf",
    "MK_De07_Lop6": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Minh Khanh - 6A3\De 07 - Lop 6.pdf",
    "MK_De08_Lop6": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Minh Khanh - 6A3\De 08 - Lop 6.pdf",
    "MK_De01_Lop8": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Minh Khanh - 6A3\De 01 - Lop 8.pdf",
    "MK_De02_Lop8": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Minh Khanh - 6A3\De 02 - Lop 8.pdf",
    "MK_De03_Lop8": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Minh Khanh - 6A3\De 03 - Lop 8.pdf",
    "MK_De04_Lop8": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Minh Khanh - 6A3\De 04 - Lop 8.pdf",
    "MK_De05_Lop8": r"H:\My Drive\Claude Code\Kiem Tra Online SuKhoiKem\Minh Khanh - 6A3\De 05 - Lop 8.pdf",
}

print("="*80)
for key, path in files.items():
    print(f"\n{'='*80}")
    print(f"FILE: {key}")
    print(f"PATH: {path}")
    print(f"{'='*80}")
    ext = os.path.splitext(path)[1].lower()
    if ext == '.pdf':
        text = read_pdf(path)
    elif ext in ('.docx', '.doc'):
        text = read_docx(path)
    else:
        text = f"[Unknown extension: {ext}]"
    print(text[:8000])  # Print first 8000 chars per file
    print()
